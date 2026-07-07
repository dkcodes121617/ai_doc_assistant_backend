import io
import base64
import logging
import zipfile
import pdfplumber
import pypdfium2 as pdfium
from PIL import Image
import docx
from google.genai import types
from app.core.clients import get_genai_client
from app.core.errors import is_quota_error

logger = logging.getLogger(__name__)

# Minimum characters on a page to consider text-based (below = image-only)
_MIN_TEXT_CHARS = 30


def _page_to_base64_png(pdf_bytes: bytes, page_index: int) -> str:
    """Render a single PDF page to a base64-encoded PNG using pypdfium2."""
    doc = pdfium.PdfDocument(pdf_bytes)
    page = doc[page_index]
    bitmap = page.render(scale=2.0)          # 2× scale → ~144 dpi, good for OCR
    pil_image = bitmap.to_pil()
    buf = io.BytesIO()
    pil_image.save(buf, format="PNG", optimize=True)
    buf.seek(0)
    return base64.b64encode(buf.read()).decode("utf-8")


def _ocr_page_with_gemini(pdf_bytes: bytes, page_index: int, page_num: int) -> str:
    """Send a rendered PDF page image to Gemini Vision and return extracted text."""
    try:
        b64 = _page_to_base64_png(pdf_bytes, page_index)
        client = get_genai_client()
        response = client.models.generate_content(
            model="gemini-2.5-flash",
            contents=[
                types.Part.from_bytes(
                    data=base64.b64decode(b64),
                    mime_type="image/png",
                ),
                (
                    f"This is page {page_num} of a PDF document. "
                    "Extract ALL text exactly as it appears, preserving structure "
                    "(paragraphs, bullet points, headings, tables). "
                    "Return only the extracted text, no commentary."
                ),
            ],
        )
        return (response.text or "").strip()
    except Exception as e:
        logger.warning("[OCR] Gemini Vision failed for page %s: %s", page_num, e)
        # Out-of-credits / rate-limit should surface to the user, not be swallowed.
        if is_quota_error(e):
            raise
        return ""


def extract_text_from_pdf(file_bytes: bytes):
    """Extract text from PDF. Falls back to Gemini Vision OCR for image-only pages. Yields progress."""
    pages_data = []
    yield {"type": "progress", "status": "Opening PDF document...", "percent": 5}

    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        total = len(pdf.pages)
        for i, page in enumerate(pdf.pages):
            text = (page.extract_text() or "").strip()

            if len(text) < _MIN_TEXT_CHARS:
                # Image-only or scanned page — use Gemini Vision OCR
                yield {"type": "progress", "status": f"Detected image on page {i + 1}, running Vision OCR...", "percent": int(5 + (i/total)*40)}
                logger.info("[OCR] Page %s has <%s chars, using Vision OCR", i + 1, _MIN_TEXT_CHARS)
                text = _ocr_page_with_gemini(file_bytes, i, i + 1)
            else:
                yield {"type": "progress", "status": f"Extracting text from page {i + 1} of {total}...", "percent": int(5 + (i/total)*40)}

            if text:
                pages_data.append({
                    "page_number": i + 1,
                    "text": text,
                })

    yield {"type": "result", "data": pages_data}


def extract_text_from_txt(file_bytes: bytes):
    yield {"type": "progress", "status": "Reading text file...", "percent": 20}
    text = file_bytes.decode("utf-8", errors="replace").strip()
    pages_data = [{"page_number": 1, "text": text}] if text else []
    yield {"type": "result", "data": pages_data}


def extract_text_from_docx(file_bytes: bytes):
    """Extract text from DOCX, grouping paragraphs into synthetic pages."""
    yield {"type": "progress", "status": "Opening Word document...", "percent": 5}
    doc = docx.Document(io.BytesIO(file_bytes))
    pages_data: list[dict] = []
    current_lines: list[str] = []
    current_len = 0
    page_num = 1
    PAGE_CHAR_LIMIT = 3000

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        current_lines.append(text)
        current_len += len(text)
        if current_len >= PAGE_CHAR_LIMIT:
            pages_data.append({
                "page_number": page_num,
                "text": "\n".join(current_lines),
            })
            page_num += 1
            current_lines = []
            current_len = 0

    if current_lines:
        pages_data.append({
            "page_number": page_num,
            "text": "\n".join(current_lines),
        })

    # If DOCX has no text, try to find and OCR embedded images
    if not pages_data:
        yield {"type": "progress", "status": "No text found. Extracting images for Vision OCR...", "percent": 20}
        try:
            with zipfile.ZipFile(io.BytesIO(file_bytes)) as z:
                image_texts = []
                client = get_genai_client()
                media_files = [n for n in z.namelist() if n.startswith("word/media/") and n.lower().endswith((".png", ".jpg", ".jpeg"))]
                total_imgs = len(media_files)
                for i, name in enumerate(media_files):
                    yield {"type": "progress", "status": f"Running OCR on image {i+1} of {total_imgs}...", "percent": int(20 + (i/total_imgs)*25)}
                    img_bytes = z.read(name)
                    mime = "image/jpeg" if name.lower().endswith((".jpg", ".jpeg")) else "image/png"
                    try:
                        response = client.models.generate_content(
                            model="gemini-2.5-flash",
                            contents=[
                                types.Part.from_bytes(data=img_bytes, mime_type=mime),
                                "Extract ALL text from this image exactly as it appears. Return only the extracted text."
                            ]
                        )
                        if response.text:
                            image_texts.append(response.text.strip())
                    except Exception as e:
                        if is_quota_error(e):
                            raise
                        logger.warning("Vision OCR failed on embedded image: %s", e)
                if image_texts:
                    pages_data.append({
                        "page_number": 1,
                        "text": "\n\n".join(image_texts)
                    })
        except Exception as e:
            # Let credit/rate-limit errors bubble up; swallow benign parse issues.
            if is_quota_error(e):
                raise
            logger.warning("Failed to process DOCX images: %s", e)

    yield {"type": "result", "data": pages_data}
